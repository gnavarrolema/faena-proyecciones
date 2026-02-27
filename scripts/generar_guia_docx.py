"""
Script para generar la Guía de Usuario de ProyecFaena en formato DOCX.
Ejecutar: python scripts/generar_guia_docx.py
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from pathlib import Path
import datetime


# ─── Colores ────────────────────────────────────────────────────────────────────
PRIMARY = RGBColor(0x1A, 0x56, 0x32)       # verde oscuro
PRIMARY_LIGHT = RGBColor(0x22, 0x8B, 0x22) # verde
ACCENT = RGBColor(0x0E, 0x7A, 0xC4)        # azul
TEXT_DARK = RGBColor(0x1E, 0x29, 0x3B)
TEXT_LIGHT = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "1A5632"
TABLE_ALT_BG = "F1F5F9"
ORANGE = RGBColor(0xEA, 0x58, 0x0C)
RED = RGBColor(0xEF, 0x44, 0x44)
GREEN = RGBColor(0x16, 0xA3, 0x4A)


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header row and alternating row colors."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            run.font.color.rgb = TEXT_DARK
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    # Column widths
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)

    doc.add_paragraph()  # spacing
    return table


def add_tip(doc, text, prefix="Tip"):
    """Add a highlighted tip/note paragraph."""
    p = doc.add_paragraph()
    run_prefix = p.add_run(f"  {prefix}: ")
    run_prefix.bold = True
    run_prefix.font.color.rgb = ACCENT
    run_prefix.font.size = Pt(9.5)
    run_text = p.add_run(text)
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = TEXT_LIGHT
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)


def add_note(doc, text):
    add_tip(doc, text, prefix="Nota")


def add_important(doc, text):
    add_tip(doc, text, prefix="Importante")


def add_warning(doc, text):
    p = doc.add_paragraph()
    run_prefix = p.add_run("  ⚠ Precaución: ")
    run_prefix.bold = True
    run_prefix.font.color.rgb = ORANGE
    run_prefix.font.size = Pt(9.5)
    run_text = p.add_run(text)
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = TEXT_LIGHT
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)


def add_numbered_steps(doc, steps):
    """Add numbered steps."""
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"{i}. ")
        run_num.bold = True
        run_num.font.size = Pt(10)
        run_text = p.add_run(step)
        run_text.font.size = Pt(10)
        pf = p.paragraph_format
        pf.left_indent = Cm(1.5)
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        run = p.add_run(item)
        run.font.size = Pt(10)


def build_document():
    doc = Document()

    # ── Page setup ───────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Styles ───────────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = TEXT_DARK
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'
        h_style.font.color.rgb = PRIMARY
        if level == 1:
            h_style.font.size = Pt(20)
            h_style.paragraph_format.space_before = Pt(24)
            h_style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            h_style.font.size = Pt(15)
            h_style.paragraph_format.space_before = Pt(18)
            h_style.paragraph_format.space_after = Pt(8)
        else:
            h_style.font.size = Pt(12)
            h_style.paragraph_format.space_before = Pt(14)
            h_style.paragraph_format.space_after = Pt(6)

    # ═════════════════════════════════════════════════════════════════════════
    # PORTADA
    # ═════════════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GUÍA DE USUARIO")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ProyecFaena")
    run.font.size = Pt(24)
    run.font.color.rgb = ACCENT

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Proyección de Faena Avícola")
    run.font.size = Pt(14)
    run.font.color.rgb = TEXT_LIGHT

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Versión 1.0 — {datetime.date.today().strftime('%B %Y').title()}")
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_LIGHT

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # ÍNDICE
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('Índice', level=1)

    toc_items = [
        "1. ¿Qué es ProyecFaena?",
        "2. Requisitos previos",
        "3. Acceso al sistema",
        "4. Pantalla principal — Navegación",
        "5. Cargar Oferta",
        "6. Ver Oferta",
        "7. Proyección de Faena",
        "8. Resumen Semanal",
        "9. Parámetros de Cálculo",
        "10. Exportación a PDF",
        "11. Semáforo de edades",
        "12. Flujo de trabajo recomendado",
        "13. Glosario de términos",
        "14. Preguntas frecuentes",
        "15. Solución de problemas",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 1. ¿Qué es ProyecFaena?
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('1. ¿Qué es ProyecFaena?', level=1)

    doc.add_paragraph(
        'ProyecFaena es una aplicación web de planificación de faena avícola. Permite:'
    )

    add_bullet_list(doc, [
        'Cargar la oferta de granjas de engorde recibida cada jueves (y ajustarla con la oferta del martes).',
        'Generar automáticamente la proyección semanal de retiro de pollos, distribuyéndolos de lunes a sábado según peso, edad y demanda (objetivo ~30.000 a 35.000 pollos/día).',
        'Ajustar manualmente la distribución moviendo lotes entre días.',
        'Visualizar indicadores clave como peso promedio, calibre, cajas producidas y diferencia de edad ideal.',
        'Exportar reportes en formato PDF.',
    ])

    doc.add_paragraph(
        'La aplicación automatiza los cálculos que tradicionalmente se realizaban en una hoja de Excel, '
        'ofreciendo precisión, velocidad y trazabilidad.'
    )

    # ═════════════════════════════════════════════════════════════════════════
    # 2. Requisitos previos
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('2. Requisitos previos', level=1)

    add_styled_table(doc,
        ['Requisito', 'Detalle'],
        [
            ['Navegador', 'Google Chrome, Microsoft Edge, Firefox o Safari (versión actual)'],
            ['Conexión', 'Acceso a la red donde está desplegado el servidor'],
            ['Archivo Excel', 'Formato .xlsx o .xls con la oferta de granjas (ver sección 5.1)'],
            ['Credenciales', 'Usuario y contraseña proporcionados por el administrador'],
        ],
        col_widths=[4, 13]
    )

    # ═════════════════════════════════════════════════════════════════════════
    # 3. Acceso al sistema
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('3. Acceso al sistema', level=1)

    doc.add_heading('3.1 Página de inicio', level=2)
    doc.add_paragraph(
        'Al ingresar a la URL de la aplicación, verá la Landing Page con información general del sistema. '
        'Desde aquí puede hacer clic en:'
    )
    add_bullet_list(doc, [
        '"Iniciar Sesión" (esquina superior derecha)',
        '"Comenzar Ahora" (botón central)',
    ])
    doc.add_paragraph('Ambos lo llevan a la pantalla de login.')

    doc.add_heading('3.2 Iniciar sesión', level=2)
    add_numbered_steps(doc, [
        'Ingrese su Usuario en el primer campo.',
        'Ingrese su Contraseña en el segundo campo.',
        'Haga clic en "Ingresar".',
    ])
    doc.add_paragraph(
        'Si los datos son correctos, verá el mensaje "¡Bienvenido!" y será redirigido a la pantalla principal.'
    )
    add_note(doc, 'La sesión permanece activa por 7 días. Si el token expira, será redirigido automáticamente al login.')
    add_note(doc, 'Credenciales por defecto (desarrollo): Usuario: admin / Contraseña: admin123. En producción, el administrador le proporcionará sus credenciales.')

    # ═════════════════════════════════════════════════════════════════════════
    # 4. Pantalla principal
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('4. Pantalla principal — Navegación', level=1)

    doc.add_paragraph(
        'Una vez autenticado, la pantalla principal muestra una barra de navegación superior con 5 pestañas:'
    )

    add_styled_table(doc,
        ['Pestaña', 'Ícono', 'Función'],
        [
            ['Cargar Oferta', '📁', 'Subir el archivo Excel con la oferta de granjas'],
            ['Oferta', '📋', 'Ver los lotes cargados y generar la proyección'],
            ['Proyección', '📊', 'Ver y ajustar la distribución diaria de lotes'],
            ['Resumen', '📈', 'Dashboard con totales diarios y semanales'],
            ['Parámetros', '⚙️', 'Configurar los valores de cálculo'],
        ],
        col_widths=[3.5, 1.5, 12]
    )

    doc.add_paragraph('En el extremo derecho de la barra hay un botón "Salir" para cerrar sesión.')

    add_tip(doc, 'Al ingresar, la aplicación detecta automáticamente si ya hay datos cargados y lo lleva a la pestaña más relevante (Proyección si ya existe, Oferta si hay datos, o Cargar Oferta si está vacío).')

    # ═════════════════════════════════════════════════════════════════════════
    # 5. Cargar Oferta
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('5. Cargar Oferta', level=1)

    doc.add_paragraph(
        'Esta pestaña permite subir el archivo Excel con la oferta del jueves de granjas de engorde.'
    )

    doc.add_heading('5.1 Formato del archivo Excel', level=2)
    doc.add_paragraph('El archivo debe ser .xlsx o .xls con las siguientes columnas:')

    add_styled_table(doc,
        ['Columna', 'Campo', 'Ejemplo'],
        [
            ['A', 'Fecha de Peso', '12/2/2026'],
            ['B', 'Granja', 'LOS REMANSOS'],
            ['C', 'Galpón', '5'],
            ['D', 'Núcleo', '1'],
            ['E', 'Cantidad', '4.370'],
            ['F', 'Sexo (M/H)', 'H'],
            ['G', 'Edad Proyectada', '42'],
            ['H', 'Peso Muestreo Proyectado', '2,78'],
            ['I', 'Ganancia Diaria', '0,090'],
            ['J', 'Días Proyectados', '0'],
            ['K', 'Edad Real', '42'],
            ['L', 'Peso Muestreo Real', '2,78'],
            ['N', 'Fecha de Ingreso', '31/12/2025'],
        ],
        col_widths=[2, 5, 3]
    )

    add_important(doc, 'Respete el orden de columnas. El sistema parsea automáticamente los datos basándose en esta estructura.')

    doc.add_heading('5.2 Pasos para cargar la oferta', level=2)
    add_numbered_steps(doc, [
        'Haga clic en la zona de carga (cuadro punteado) o arrastre el archivo directamente sobre ella.',
        'Verifique que aparezca el nombre del archivo seleccionado.',
        'Haga clic en "Cargar y Procesar".',
        'Espere a que el sistema procese el archivo. Al finalizar, será redirigido automáticamente a la pestaña Oferta.',
    ])

    add_tip(doc, 'Si seleccionó un archivo incorrecto, use el botón "Limpiar" para descartarlo y seleccionar otro.')

    doc.add_heading('5.3 Advertencia de sobreescritura', level=2)
    doc.add_paragraph(
        'Si ya existen datos cargados en el sistema (oferta o proyección previas), aparecerá un cuadro de advertencia '
        'naranja indicando que la carga reemplazará completamente los datos actuales.'
    )
    doc.add_paragraph('Para confirmar:')
    add_numbered_steps(doc, [
        'Marque la casilla "Entiendo que los datos actuales serán reemplazados y deseo continuar".',
        'Se habilitará el botón "Cargar y Procesar".',
    ])

    add_tip(doc, 'Si solo desea actualizar los datos de peso/edad sin perder la planificación existente, use la opción "Ajuste Martes" desde la pestaña Proyección (ver sección 7.2).')

    # ═════════════════════════════════════════════════════════════════════════
    # 6. Ver Oferta
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('6. Ver Oferta', level=1)

    doc.add_paragraph(
        'Una vez cargado el archivo, esta pestaña muestra toda la información de los lotes importados.'
    )

    doc.add_heading('6.1 Resumen estadístico', level=2)
    doc.add_paragraph('En la parte superior se muestran tres tarjetas con:')
    add_bullet_list(doc, [
        'Total Lotes: Cantidad de lotes (galpones) cargados.',
        'Total Pollos: Suma total de pollos en todos los lotes.',
        'Granjas: Cantidad de granjas distintas en la oferta.',
    ])

    doc.add_heading('6.2 Generar Proyección', level=2)
    doc.add_paragraph('Este es el paso clave. Complete los campos:')

    add_styled_table(doc,
        ['Campo', 'Descripción', 'Ejemplo'],
        [
            ['Fecha Inicio Semana (Lunes)', 'El lunes de la semana a planificar', '2026-03-02'],
            ['Pollos por Día (objetivo)', 'Cuántos pollos desea faenar por día', '30000'],
            ['Días de Faena', 'Cuántos días trabajar (5 o 6)', '6 (Lunes a Sábado)'],
        ],
        col_widths=[5, 7, 4]
    )

    doc.add_paragraph('Luego haga clic en "Generar Proyección Automática". El sistema:')
    add_numbered_steps(doc, [
        'Calcula edad y peso proyectado de cada lote para cada día de la semana.',
        'Filtra los lotes que están fuera de rango de edad/peso permitido.',
        'Prioriza lotes según edad ideal y los distribuye equilibradamente entre los días.',
        'Respeta el tope de pollos por día configurado.',
    ])
    doc.add_paragraph('Al completar, será redirigido a la pestaña Proyección.')

    doc.add_heading('6.3 Resumen por Granja', level=2)
    doc.add_paragraph(
        'Tabla que muestra cada granja con la cantidad de lotes y pollos totales. '
        'Incluye un botón "Descargar PDF" para exportar este resumen.'
    )

    doc.add_heading('6.4 Tabla de Oferta Completa', level=2)
    doc.add_paragraph(
        'Tabla con todos los lotes cargados mostrando todos sus datos: fecha de peso, granja, galpón, '
        'núcleo, cantidad, sexo, edad, peso, ganancia diaria, etc. Es de solo lectura y permite verificar '
        'que los datos se importaron correctamente.'
    )

    # ═════════════════════════════════════════════════════════════════════════
    # 7. Proyección de Faena
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('7. Proyección de Faena', level=1)

    doc.add_paragraph(
        'Esta es la pestaña central de la aplicación. Aquí se visualiza y ajusta la distribución de lotes por día.'
    )

    doc.add_heading('7.1 Indicadores semanales', level=2)
    doc.add_paragraph('En la parte superior se muestran 4 tarjetas:')

    add_styled_table(doc,
        ['Indicador', 'Qué muestra'],
        [
            ['Total Pollos Semana', 'Suma de todos los pollos asignados en la semana'],
            ['Promedio Edad Semana', 'Edad promedio de retiro de todos los lotes'],
            ['Cajas Semanales', 'Total de cajas de 20 kg producidas en la semana'],
            ['Sofía (Total - 10.000)', 'Pollos semanales menos el descuento de Sofía (configurable)'],
        ],
        col_widths=[5, 12]
    )

    doc.add_heading('7.2 Ajuste con Oferta del Martes', level=2)
    doc.add_paragraph(
        'Esta sección (colapsable) permite actualizar la proyección existente con datos más frescos de la '
        'oferta del martes sin perder la distribución de días ya planificada.'
    )

    p = doc.add_paragraph()
    run = p.add_run('¿Cuándo usarlo? ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(
        'El martes llega una nueva oferta con datos de peso y edad actualizados. Este ajuste:'
    )
    run2.font.size = Pt(10)

    add_bullet_list(doc, [
        'Actualiza los datos (peso, edad, ganancia) de lotes ya existentes.',
        'Agrega nuevos lotes si hay capacidad disponible en algún día.',
        'Mantiene la distribución de días que ya fue planificada.',
        'Alerta si algún lote existente queda fuera de rango tras la actualización.',
    ])

    doc.add_paragraph('Pasos:')
    add_numbered_steps(doc, [
        'Haga clic en "Ajustar con Oferta del Martes" para expandir la sección.',
        'Seleccione el archivo Excel de la oferta del martes.',
        'Haga clic en "Aplicar Ajuste".',
        'Revise el resumen del ajuste que aparece.',
    ])

    doc.add_paragraph('El resumen indica:')
    add_bullet_list(doc, [
        '✅ Lotes actualizados (datos cambiados)',
        '✅ Lotes nuevos asignados',
        'ℹ️ Lotes nuevos sin capacidad',
        '⚠️ Lotes existentes que ahora quedaron fuera de rango',
        '⚠️ Lotes no encontrados en la nueva oferta',
    ])

    doc.add_heading('7.3 Lotes no asignados', level=2)
    doc.add_paragraph(
        'Si el sistema no pudo asignar algunos lotes por exceso de capacidad diaria, aparece una sección '
        'amarilla listando esos lotes con: Granja, Galpón, Núcleo, Cantidad de pollos, Días elegibles '
        '(en qué días podría haber ido) y Motivo (ej: "Tope diario alcanzado").'
    )
    add_tip(doc, 'Considere ajustar los parámetros de pollos por día o mover lotes manualmente para hacer espacio.')

    doc.add_heading('7.4 Lotes fuera de rango', level=2)
    doc.add_paragraph(
        'Los lotes que no cumplen los requisitos de edad o peso mínimo/máximo para ningún día de la semana '
        'aparecen en una sección roja. Puede expandir cada lote para ver el detalle día por día (edad y peso '
        'proyectado, y la razón por la que no califica).'
    )

    doc.add_heading('7.5 Vista por Día (Cards)', level=2)
    doc.add_paragraph(
        'La vista predeterminada muestra una grilla tipo Kanban con una columna por cada día de faena '
        '(Lunes a Sábado). Cada columna contiene:'
    )

    add_bullet_list(doc, [
        'Encabezado: Nombre del día y total de pollos.',
        'Tarjetas de lotes con: Granja y Galpón, Sexo (M/H/-), Cantidad de pollos, Edad al momento del retiro, '
        'Peso vivo (kg), Diferencia de edad ideal (con código de color), Peso faenado, Cajas producidas, '
        'y botones "Mover" y "Eliminar".',
        'Resumen inferior: Peso promedio, diferencia de edad promedio y cajas del día.',
    ])

    doc.add_heading('7.6 Vista Tabla', level=2)
    doc.add_paragraph(
        'Alterne a esta vista haciendo clic en "Vista Tabla". Muestra todos los lotes en una tabla única '
        'con filas agrupadas por día. Incluye subtotales por cada día con los promedios ponderados.'
    )

    doc.add_heading('7.7 Mover un lote a otro día', level=2)
    add_numbered_steps(doc, [
        'En la tarjeta o fila del lote, haga clic en "Mover".',
        'Se abrirá un diálogo mostrando los demás días disponibles con su fecha y total de pollos actual.',
        'Haga clic en el día destino deseado.',
        'El lote se recalcula automáticamente para la nueva fecha y se actualiza toda la proyección.',
    ])

    doc.add_heading('7.8 Eliminar un lote', level=2)
    add_numbered_steps(doc, [
        'Haga clic en "Eliminar" (o el ícono ✕ en la vista tabla).',
        'Confirme en el diálogo de confirmación.',
        'El lote se retira de la proyección y los totales se recalculan.',
    ])
    add_warning(doc, 'Eliminar un lote no lo devuelve a la oferta. Si necesita recuperarlo, regenere la proyección desde la pestaña Oferta.')

    # ═════════════════════════════════════════════════════════════════════════
    # 8. Resumen Semanal
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('8. Resumen Semanal', level=1)

    doc.add_paragraph('Esta pestaña ofrece una visión consolidada de la semana planificada.')

    doc.add_heading('8.1 Resumen Diario', level=2)
    doc.add_paragraph('Tabla con una fila por día que muestra:')

    add_styled_table(doc,
        ['Columna', 'Descripción'],
        [
            ['Día', 'Lunes a Sábado'],
            ['Fecha', 'Fecha calendario del día'],
            ['Pollos', 'Total de pollos del día'],
            ['Lotes', 'Cantidad de lotes asignados'],
            ['Peso Prom.', 'Peso vivo promedio ponderado (kg)'],
            ['Dif. Edad Prom.', 'Diferencia de edad promedio vs. ideal'],
            ['Calibre Prom.', 'Calibre promedio (pollos/caja) ponderado'],
            ['Cajas', 'Cajas producidas en el día'],
        ],
        col_widths=[4, 13]
    )

    doc.add_paragraph('La fila TOTAL SEMANA suma los pollos y cajas de toda la semana.')

    doc.add_heading('8.2 Distribución por Granja', level=2)
    doc.add_paragraph(
        'Tabla cruzada que muestra cuántos pollos de cada granja se procesan en cada día. '
        'Permite visualizar rápidamente la distribución y detectar si alguna granja se concentra en un solo día.'
    )

    doc.add_heading('8.3 Cobertura de la Oferta', level=2)
    doc.add_paragraph('Si existen lotes fuera de rango o no asignados, aparece una sección adicional que muestra:')
    add_bullet_list(doc, [
        'Total Ofertados: Todos los pollos de la oferta original.',
        'Asignados (%): Cuántos se incorporaron a la proyección.',
        'Fuera de Rango: Pollos que no cumplen edad/peso.',
        'Exceso de Capacidad: Pollos elegibles pero sin espacio por tope diario.',
    ])
    doc.add_paragraph('Incluye un botón "Descargar PDF" para exportar este resumen.')

    # ═════════════════════════════════════════════════════════════════════════
    # 9. Parámetros de Cálculo
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('9. Parámetros de Cálculo', level=1)

    doc.add_paragraph(
        'En esta pestaña puede configurar todos los valores que afectan los cálculos de la proyección.'
    )

    doc.add_heading('9.1 Ganancia de Peso', level=3)
    add_styled_table(doc,
        ['Parámetro', 'Valor por defecto', 'Descripción'],
        [
            ['Ganancia diaria machos', '0.090 kg', 'Incremento de peso diario de pollos machos'],
            ['Ganancia diaria hembras', '0.079 kg', 'Incremento de peso diario de pollos hembras'],
            ['Factor medio día', '0.5', 'Factor aplicado a la ganancia del último día'],
        ],
        col_widths=[5, 3.5, 8.5]
    )

    doc.add_heading('9.2 Rendimiento', level=3)
    add_styled_table(doc,
        ['Parámetro', 'Valor por defecto', 'Descripción'],
        [
            ['Rendimiento canal', '0.87 (87%)', 'Proporción del peso vivo que queda tras la faena'],
            ['Kg por caja', '20.0 kg', 'Peso estándar por caja'],
            ['Descuento sin sexar', '0.04 (4%)', 'Penalización aplicada a pollos sin sexar'],
        ],
        col_widths=[5, 3.5, 8.5]
    )

    doc.add_heading('9.3 Edades Ideales', level=3)
    add_styled_table(doc,
        ['Parámetro', 'Valor por defecto', 'Descripción'],
        [
            ['Edad ideal machos', '40 días', 'Edad óptima de retiro para machos'],
            ['Edad ideal hembras', '44 días', 'Edad óptima de retiro para hembras'],
            ['Edad ideal sin sexar', '42 días', 'Edad óptima de retiro para lotes sin sexar'],
            ['Edad mínima faena', '—', 'Edad mínima para que un lote sea elegible'],
            ['Edad máxima faena', '—', 'Edad máxima para que un lote sea elegible'],
        ],
        col_widths=[5, 3.5, 8.5]
    )

    doc.add_heading('9.4 Rango de Peso Faena', level=3)
    add_styled_table(doc,
        ['Parámetro', 'Descripción'],
        [
            ['Peso mínimo faena', 'Peso vivo mínimo para que un lote sea elegible'],
            ['Peso máximo faena', 'Peso vivo máximo para que un lote sea elegible'],
        ],
        col_widths=[5, 12]
    )

    doc.add_heading('9.5 Producción', level=3)
    add_styled_table(doc,
        ['Parámetro', 'Valor por defecto', 'Descripción'],
        [
            ['Pollos diarios mín.', '30.000', 'Objetivo mínimo de pollos por día'],
            ['Pollos diarios máx.', '35.000', 'Objetivo máximo de pollos por día'],
            ['Descuento Sofía', '10.000', 'Constante que se resta del total semanal'],
        ],
        col_widths=[5, 3.5, 8.5]
    )

    doc.add_heading('9.6 Guardar cambios', level=2)
    add_numbered_steps(doc, [
        'Modifique los valores deseados.',
        'Haga clic en "Guardar" (esquina superior derecha).',
        'Aparecerá un mensaje verde de confirmación.',
    ])

    add_important(doc, 'Los cambios en parámetros no se aplican retroactivamente a la proyección existente. Debe regenerar la proyección desde la pestaña Oferta para que los nuevos valores tomen efecto.')

    doc.add_paragraph('También puede descargar los parámetros actuales en PDF con el botón "Descargar PDF".')

    # ═════════════════════════════════════════════════════════════════════════
    # 10. Exportación a PDF
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('10. Exportación a PDF', level=1)

    doc.add_paragraph('La aplicación permite exportar varios reportes en formato PDF:')

    add_styled_table(doc,
        ['Reporte', 'Desde dónde', 'Contenido'],
        [
            ['Oferta', 'Pestaña Oferta → "Descargar PDF"', 'Resumen por granja + tabla completa de lotes'],
            ['Proyección', 'Pestaña Proyección → "Descargar PDF"', 'Distribución detallada por día con indicadores'],
            ['Resumen', 'Pestaña Resumen → "Descargar PDF"', 'Resumen diario, por granja y cobertura'],
            ['Parámetros', 'Pestaña Parámetros → "Descargar PDF"', 'Configuración actual de todos los parámetros'],
        ],
        col_widths=[3, 6, 8]
    )

    doc.add_paragraph(
        'Los PDFs se generan al instante y se descargan directamente al navegador con un nombre que '
        'incluye la fecha (ej: oferta-2026-02-27.pdf).'
    )

    # ═════════════════════════════════════════════════════════════════════════
    # 11. Semáforo de edades
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('11. Semáforo de edades', level=1)

    doc.add_paragraph(
        'En las vistas de proyección, la diferencia de edad ideal se muestra con un código de colores '
        'para identificar rápidamente si un lote está en su punto óptimo:'
    )

    add_styled_table(doc,
        ['Color', 'Rango', 'Significado'],
        [
            ['🟢 Verde', '-1 a +1 días', 'Edad óptima — el lote está dentro del rango ideal'],
            ['🟠 Naranja', '-3 a -2 o +2 a +3 días', 'Atención — el lote se aleja ligeramente del ideal'],
            ['🔴 Rojo', 'Menor a -3 o mayor a +3 días', 'Alerta — significativamente fuera de la edad ideal'],
        ],
        col_widths=[3, 5, 9]
    )

    doc.add_paragraph('La diferencia de edad se calcula como:')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Diferencia = Edad al retiro − Edad ideal según sexo')
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT

    doc.add_paragraph()
    doc.add_paragraph('Donde:')
    add_bullet_list(doc, [
        'Machos: edad ideal = 40 días',
        'Hembras: edad ideal = 44 días',
        'Sin sexar: edad ideal = 42 días',
    ])

    # ═════════════════════════════════════════════════════════════════════════
    # 12. Flujo de trabajo recomendado
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('12. Flujo de trabajo recomendado', level=1)

    doc.add_paragraph('A continuación se describe el flujo semanal típico de uso de la aplicación.')

    doc.add_heading('Jueves — Carga inicial', level=2)
    add_numbered_steps(doc, [
        'Recibir el Excel de oferta de granjas (OFERTA JUEV).',
        'Iniciar sesión en ProyecFaena.',
        'Ir a "Cargar Oferta" → subir el archivo Excel.',
        'Ir a "Oferta" → verificar los datos importados.',
        'Configurar la fecha del próximo lunes, pollos/día y días de faena.',
        'Clic en "Generar Proyección Automática".',
        'Ir a "Proyección" → revisar la distribución.',
        'Mover lotes entre días si es necesario.',
        'Ir a "Resumen" → validar los totales.',
        'Exportar a PDF para compartir con el equipo.',
    ])

    doc.add_heading('Martes — Ajuste con datos actualizados', level=2)
    add_numbered_steps(doc, [
        'Recibir el Excel de oferta actualizada (OFERTA MART).',
        'Ir a "Proyección" → abrir "Ajustar con Oferta del Martes".',
        'Subir el archivo y aplicar ajuste.',
        'Revisar el resumen de cambios.',
        'Verificar si hay lotes que ahora quedan fuera de rango.',
        'Ajustar manualmente si es necesario.',
        'Exportar el PDF actualizado.',
    ])

    doc.add_heading('Si necesita rehacer todo desde cero', level=2)
    add_numbered_steps(doc, [
        'Ir a "Cargar Oferta".',
        'Confirmar la sobreescritura.',
        'Subir nuevo archivo.',
        'Regenerar la proyección.',
    ])

    # ═════════════════════════════════════════════════════════════════════════
    # 13. Glosario
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('13. Glosario de términos', level=1)

    add_styled_table(doc,
        ['Término', 'Definición'],
        [
            ['Lote', 'Un grupo de pollos de un mismo galpón, granja y núcleo'],
            ['Oferta', 'Conjunto de lotes disponibles para faena, informado por las granjas'],
            ['Proyección', 'Planificación de qué lotes se retiran cada día de la semana'],
            ['Faena', 'Proceso industrial de sacrificio y procesamiento de los pollos'],
            ['Peso vivo', 'Peso estimado del pollo al momento del retiro en la granja'],
            ['Peso faenado', 'Peso del pollo después de la faena (= peso vivo × rendimiento canal)'],
            ['Rendimiento canal', 'Porcentaje del peso vivo que se conserva tras la faena (87%)'],
            ['Calibre', 'Cantidad de pollos que caben en una caja de 20 kg'],
            ['Cajas', 'Unidad de producción; cada caja contiene 20 kg de pollo faenado'],
            ['Ganancia diaria', 'Cuántos kg de peso gana un pollo por día'],
            ['Edad ideal', 'Edad óptima de retiro según el sexo del pollo'],
            ['Diferencia de edad', 'Días de más o de menos respecto a la edad ideal'],
            ['Descuento Sofía', 'Constante (10.000 pollos) que se resta del total semanal'],
            ['Descuento sin sexar', 'Penalización del 4% en peso para pollos sin sexo determinado'],
            ['Galpón', 'Estructura dentro de una granja donde se crían los pollos'],
            ['Núcleo', 'Subdivisión dentro de una granja (agrupación de galpones)'],
        ],
        col_widths=[4, 13]
    )

    # ═════════════════════════════════════════════════════════════════════════
    # 14. Preguntas frecuentes
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('14. Preguntas frecuentes', level=1)

    faqs = [
        ('¿Puedo cargar la oferta más de una vez?',
         'Sí. Cada vez que carga una nueva oferta desde "Cargar Oferta", reemplaza completamente la oferta anterior y la proyección. Si solo desea actualizar datos de peso/edad sin perder la planificación, use "Ajustar con Oferta del Martes" en la pestaña Proyección.'),

        ('¿Qué pasa si el archivo Excel tiene un formato diferente?',
         'El sistema esperará las columnas en el orden indicado en la sección 5.1. Si el formato difiere, se producirá un error con un mensaje descriptivo.'),

        ('¿Los cambios en Parámetros afectan la proyección existente?',
         'No de forma automática. Los parámetros se aplican al generar una nueva proyección o al aplicar un ajuste martes. Si desea que los nuevos parámetros se reflejen, debe regenerar la proyección.'),

        ('¿Puedo trabajar con 5 días de faena en vez de 6?',
         'Sí. Al generar la proyección, seleccione "5 días" en el menú desplegable. La distribución se hará de lunes a viernes.'),

        ('¿Qué significa "Sofía"?',
         'Es un indicador que resta una cantidad fija (por defecto 10.000) al total de pollos semanales. Representa un ajuste contractual o de destino específico. El valor es configurable desde Parámetros.'),

        ('¿Puedo agregar un lote manualmente?',
         'Actualmente, los lotes se agregan cargando el archivo Excel. La interfaz principal no tiene un formulario de ingreso manual individual, pero la API del backend sí soporta esta operación.'),

        ('¿Mis datos se pierden al cerrar el navegador?',
         'No. Los datos (oferta, proyección, parámetros) se almacenan en el servidor y persisten entre sesiones. Al volver a iniciar sesión, se cargarán automáticamente.'),

        ('¿Puedo usar la app desde el celular?',
         'La interfaz es responsive y funciona en dispositivos móviles, aunque la experiencia óptima es en pantallas de escritorio o tablet dado el volumen de datos tabulares.'),
    ]

    for question, answer in faqs:
        p = doc.add_paragraph()
        run_q = p.add_run(question)
        run_q.bold = True
        run_q.font.size = Pt(10.5)
        run_q.font.color.rgb = PRIMARY

        p2 = doc.add_paragraph(answer)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(10)

    # ═════════════════════════════════════════════════════════════════════════
    # 15. Solución de problemas
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading('15. Solución de problemas', level=1)

    add_styled_table(doc,
        ['Problema', 'Causa posible', 'Solución'],
        [
            ['"Usuario o contraseña incorrectos"', 'Credenciales erróneas', 'Verifique con el administrador sus datos de acceso'],
            ['"Error al procesar el archivo"', 'Formato de Excel incorrecto', 'Revise las columnas según sección 5.1'],
            ['"No hay oferta cargada"', 'No se subió archivo', 'Vaya a "Cargar Oferta" y suba el Excel'],
            ['"No hay proyección generada"', 'Falta ejecutar la generación', 'Vaya a "Oferta" → "Generar Proyección"'],
            ['Muchos lotes fuera de rango', 'Parámetros muy restrictivos', 'Ajuste rangos en "Parámetros"'],
            ['Lotes no asignados', 'Más pollos que capacidad', 'Aumente "Pollos diarios máx." o días de faena'],
            ['Página se queda cargando', 'Problema de red/servidor', 'Recargue con F5. Verifique el servidor'],
            ['Sesión expirada', 'Token JWT caducado', 'Inicie sesión nuevamente'],
            ['PDF sale en blanco', 'Navegador bloquea descargas', 'Permita descargas para este sitio'],
        ],
        col_widths=[5, 4.5, 7.5]
    )

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('¿Necesita ayuda adicional? Contacte al administrador del sistema o al equipo de desarrollo.')
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_LIGHT
    run.italic = True

    return doc


if __name__ == '__main__':
    output_path = Path(__file__).resolve().parent.parent / 'docs' / 'GUIA_DE_USUARIO.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = build_document()
    doc.save(str(output_path))
    print(f"✅ Guía generada exitosamente: {output_path}")
