"""
Script para ampliar la Guía Rápida de Usuario con las funcionalidades
de las Fases 1, 2 y 3 de integración de Producción.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX_PATH = os.path.join(BASE_DIR, "docs", "GUIA_RAPIDA_USUARIO.docx")


def add_section_title(doc, text):
    """Agrega un título de sección (estilo Normal, negrita) — igual que secciones 7-9."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p


def add_subtitle(doc, text):
    """Agrega un subtítulo (negrita, tamaño intermedio)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p


def add_normal(doc, text):
    """Agrega un párrafo normal."""
    p = doc.add_paragraph(text, style="Normal")
    return p


def add_bullet(doc, text):
    """Agrega un ítem con viñeta."""
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def main():
    doc = Document(DOCX_PATH)

    # ── Separador visual ──
    doc.add_paragraph()  # línea en blanco

    # ================================================================
    # SECCIÓN 10 — Validación de Factibilidad con Producción Propia
    # ================================================================
    add_section_title(doc, "10. Validación de Factibilidad con Producción Propia")

    add_subtitle(doc, "¿Qué es y para qué sirve?")
    add_normal(
        doc,
        "Al generar o consultar la proyección semanal, el sistema cruza automáticamente "
        "la oferta total de pollos con los datos de Cargas de Pollitos BB para determinar "
        "si la producción propia alcanza a cubrir la demanda. Si los pollitos cargados en "
        "granja (descontando mortalidad estimada) superan o igualan la oferta, se muestra "
        "un indicador positivo; de lo contrario, se alerta sobre el déficit.",
    )

    add_subtitle(doc, "¿Dónde se visualiza?")
    add_bullet(
        doc,
        "Banner en Proyección: En la pestaña Proyección, debajo de las tarjetas de "
        "estadísticas, aparece un banner de color que indica el estado de factibilidad.",
    )
    add_bullet(
        doc,
        "Tabla de Cobertura Multi-Escenario en Resumen: En la pestaña Resumen, dentro "
        "de la sección 'Referencia de Cargas en Granja', aparece la tabla "
        "'Cobertura por escenario de mortalidad' con 5 filas (una por cada tasa: "
        "4.5%, 5.0%, 5.5%, 6.0% y 6.5%).",
    )

    add_subtitle(doc, "¿Cómo interpretarlo?")
    add_bullet(
        doc,
        "Banner Verde — Producción OK: La producción propia (al peor escenario de "
        "6.5% de mortalidad) cubre la oferta total. Muestra la cantidad disponible y "
        "el porcentaje de cobertura.",
    )
    add_bullet(
        doc,
        "Banner Naranja — Déficit detectado: La oferta excede la producción propia "
        "disponible al peor escenario. Indica cuántos pollos faltan, el porcentaje "
        "de cobertura y ofrece un enlace directo para 'Agregar compra a terceros'.",
    )
    add_bullet(
        doc,
        "Tabla Multi-Escenario: Cada fila muestra la tasa de mortalidad, la cantidad "
        "de pollos disponibles, la oferta, el porcentaje de cobertura (verde si ≤100%, "
        "naranja entre 80-105%, rojo si >105%) y la diferencia (positiva = déficit en "
        "rojo, negativa = margen disponible en verde).",
    )

    # ================================================================
    # SECCIÓN 11 — Forecast de Producción (Mejor / Peor Caso)
    # ================================================================
    add_section_title(doc, "11. Forecast de Producción (Mejor / Peor Caso)")

    add_subtitle(doc, "¿Qué es y para qué sirve?")
    add_normal(
        doc,
        "El Forecast de Producción proyecta cuántos pollos estarán disponibles para "
        "faena en las próximas semanas (por defecto las 4 siguientes), basándose en "
        "los registros de cargas de pollitos y aplicando los escenarios extremos de "
        "mortalidad (4.5% mejor caso y 6.5% peor caso). Esto permite anticipar si "
        "habrá suficiente producción propia para cubrir las faenas futuras.",
    )

    add_subtitle(doc, "¿Dónde se visualiza?")
    add_bullet(
        doc,
        "Pestaña Producción (Cargas Pollitos BB): Debajo de la tabla de simulación "
        "de mortalidad aparece la sección 'Forecast de Producción' con un ícono de "
        "tendencia ascendente.",
    )

    add_subtitle(doc, "¿Cómo se interpreta la tabla?")
    add_bullet(doc, "Semana de Faena: Rango de fechas (lunes a domingo) de la semana proyectada.")
    add_bullet(doc, "Semanas Incluidas: Cantidad de semanas de carga que alimentan esa semana de faena.")
    add_bullet(
        doc,
        "Mejor Caso (azul): Pollos disponibles aplicando 4.5% de mortalidad — "
        "el escenario más optimista.",
    )
    add_bullet(
        doc,
        "Peor Caso (naranja): Pollos disponibles aplicando 6.5% de mortalidad — "
        "el escenario más conservador.",
    )
    add_bullet(doc, "Rango: Muestra el intervalo entre peor y mejor caso en formato 'X – Y'.")

    # ================================================================
    # SECCIÓN 12 — Escenarios con Tasa de Mortalidad Variable
    # ================================================================
    add_section_title(doc, "12. Escenarios con Tasa de Mortalidad Variable")

    add_subtitle(doc, "¿Qué es y para qué sirve?")
    add_normal(
        doc,
        "Al guardar un escenario de planificación, ahora es posible asociarle una tasa "
        "de mortalidad específica (de 4.5% a 6.5%). Esto permite comparar diferentes "
        "escenarios no solo por distribución de lotes, sino también por el impacto de "
        "distintas tasas de mortalidad sobre la producción propia disponible.",
    )

    add_subtitle(doc, "¿Cómo se utiliza?")
    add_bullet(
        doc,
        "Guardar escenario: En la pestaña Escenarios, al presionar 'Guardar Escenario "
        "Actual', aparece el campo 'Mortalidad % (opcional)' con las opciones: "
        "4.5% (mejor caso), 5.0%, 5.5%, 6.0% y 6.5% (peor caso). Si no se selecciona "
        "ninguna, el escenario se guarda sin análisis de producción.",
    )
    add_bullet(
        doc,
        "Badge en tarjetas: Los escenarios que tienen tasa de mortalidad asociada "
        "muestran un indicador 'Mort. X%' en su tarjeta resumen.",
    )
    add_bullet(
        doc,
        "Tabla comparativa: Al seleccionar escenarios para comparar, la tabla incluye "
        "filas adicionales de 'Tasa Mortalidad', 'Pollitos Disponibles' y "
        "'Déficit Prod.' para evaluar el impacto de la mortalidad en cada escenario.",
    )

    # ================================================================
    # SECCIÓN 13 — Tendencia de Mortalidad Observada
    # ================================================================
    add_section_title(doc, "13. Tendencia de Mortalidad Observada")

    add_subtitle(doc, "¿Qué es y para qué sirve?")
    add_normal(
        doc,
        "Esta funcionalidad retro-calcula la mortalidad real observada comparando los "
        "pollitos cargados en granja con los pollos efectivamente recibidos en planta. "
        "Permite detectar si la mortalidad real está dentro del rango esperado "
        "(4.5% – 6.5%) y observar tendencias a lo largo del tiempo.",
    )

    add_subtitle(doc, "¿Dónde se visualiza?")
    add_bullet(
        doc,
        "Pestaña Desvíos: Debajo de la sección de pesos, aparece 'Tendencia de "
        "Mortalidad Observada' con un borde de color según la tendencia general "
        "(verde = favorable, azul = normal, rojo = desfavorable).",
    )

    add_subtitle(doc, "¿Cómo se interpreta?")
    add_bullet(
        doc,
        "Resumen general: Un banner muestra el promedio de mortalidad observada, "
        "el rango, la cantidad de semanas analizadas y la tendencia.",
    )
    add_bullet(
        doc,
        "Tarjetas de estadísticas: Se muestran 3 tarjetas: Mortalidad Promedio, "
        "Rango Observado (mín % – máx %) y Semanas Analizadas.",
    )
    add_bullet(
        doc,
        "Tabla de detalle: Muestra por cada semana de carga: la fecha de carga, "
        "la fecha de faena estimada (carga + 42 días), los pollitos cargados, "
        "los pollos recibidos, el porcentaje de mortalidad observada y un estado.",
    )
    add_bullet(
        doc,
        "Estados de evaluación: Excelente (ícono verde, mortalidad ≤ 4.5%), "
        "Dentro del Rango (ícono azul, entre 4.5% y 6.5%), "
        "Por Encima (ícono rojo, mortalidad > 6.5%).",
    )
    add_bullet(
        doc,
        "Al pie de la tabla se indica el rango esperado de mortalidad "
        "(4.5% – 6.5%) según las tasas configuradas en el sistema.",
    )

    # SECCIÓN 14 — Semana 2 / Diferir Lotes
    add_section_title(doc, "14. Semana 2 — Proyección Tentativa y Diferir Lotes")
    add_normal(
        doc,
        "Cuando un feriado u otras razones comprimen la semana de faena, puede diferir "
        "lotes a la semana siguiente para visualizar cómo quedaría la planificación.",
    )
    add_subtitle(doc, "Diferir un lote")
    add_bullet(
        doc,
        "En la tarjeta del lote (Vista por Día), haga clic en el botón «S2» (ícono →).",
    )
    add_bullet(
        doc,
        "El lote se retira de Semana 1 y sus totales se recalculan al instante.",
    )
    add_bullet(
        doc,
        "Los lotes diferidos aparecen en la sección «Semana 2 — Proyección Tentativa» "
        "al pie de la pestaña Proyección.",
    )
    add_subtitle(doc, "Sección Semana 2")
    add_bullet(
        doc,
        "Badge TENTATIVA: indica que la Semana 2 es solo orientativa y de solo lectura.",
    )
    add_bullet(
        doc,
        "KPIs de S2: total pollos, promedio de edad, cajas y días de faena proyectados.",
    )
    add_bullet(
        doc,
        "Tabla de diferidos: muestra los lotes diferidos con su granja, "
        "día de origen en S1 y motivo. Incluye el botón «Restaurar» (ícono ↩).",
    )
    add_bullet(
        doc,
        "Grilla de días S2: vista Kanban de cómo quedarían los días de la semana "
        "siguiente (solo lectura, sin botones de Mover o Eliminar).",
    )
    add_bullet(
        doc,
        "Lotes fuera de rango en S2: pollos que no alcanzarán el peso/edad "
        "mínimo para la semana siguiente.",
    )
    add_subtitle(doc, "Restaurar y limpiar diferidos")
    add_bullet(
        doc,
        "Para devolver un lote a S1: clic en «Restaurar» en la tabla de diferidos. "
        "El lote se reincorpora al día con mayor déficit.",
    )
    add_bullet(
        doc,
        "Al cargar la nueva oferta la semana siguiente, use «Limpiar diferidos» "
        "para reiniciar la lista de lotes diferidos.",
    )

    # ── Guardar ──
    out_path = DOCX_PATH.replace(".docx", "_v2.docx")
    doc.save(out_path)
    print(f"✅ Documento guardado en: {out_path}")
    print("   Cierre el archivo original y renómbrelo manualmente, o copie el _v2 sobre el original.")


if __name__ == "__main__":
    main()
